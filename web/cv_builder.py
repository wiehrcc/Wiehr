import io
import os
from pathlib import Path

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.colors import Color

from django.conf import settings

CACHE_DIR = Path(settings.BASE_DIR) / "media" / "cv" / "cache"
CACHE_DIR.mkdir(parents=True, exist_ok=True)

FONT_MAP = {"bold": "Courier-Bold", "regular": "Courier"}


TEXT_FULL = Color(0, 0, 0, 1.0)
TEXT_SECONDARY = Color(0.20, 0.20, 0.20, 1.0)


def _get_styles():
    fm = FONT_MAP
    return {
        "name": ParagraphStyle("CVName", fontName=fm["bold"], fontSize=18, leading=20, textColor=TEXT_FULL, spaceAfter=2),
        "title": ParagraphStyle("CVTitle", fontName=fm["regular"], fontSize=9, leading=11.5, textColor=TEXT_FULL, spaceAfter=3),
        "contact": ParagraphStyle("CVContact", fontName=fm["regular"], fontSize=8, leading=10.5, textColor=TEXT_SECONDARY, spaceAfter=2),
        "bio": ParagraphStyle("CVBio", fontName=fm["regular"], fontSize=8, leading=11, textColor=TEXT_FULL, spaceAfter=5),
        "section": ParagraphStyle("CVSection", fontName=fm["bold"], fontSize=8.5, leading=11, textColor=TEXT_FULL, spaceBefore=3, spaceAfter=2.5),
        "role": ParagraphStyle("CVRole", fontName=fm["bold"], fontSize=9.5, leading=12, textColor=TEXT_FULL, spaceAfter=1),
        "meta": ParagraphStyle("CVMeta", fontName=fm["regular"], fontSize=8, leading=10.5, textColor=TEXT_SECONDARY, spaceAfter=1.5),
        "bullet": ParagraphStyle("CVBullet", fontName=fm["regular"], fontSize=8, leading=11, textColor=TEXT_FULL, leftIndent=8, bulletIndent=0, spaceAfter=1),
        "skill": ParagraphStyle("CVSkill", fontName=fm["regular"], fontSize=8, leading=11, textColor=TEXT_FULL, spaceAfter=1.5),
        "project_name": ParagraphStyle("CVProjectName", fontName=fm["bold"], fontSize=8.5, leading=11, textColor=TEXT_FULL, spaceAfter=1),
        "project_desc": ParagraphStyle("CVProjectDesc", fontName=fm["regular"], fontSize=8, leading=11, textColor=TEXT_FULL, spaceAfter=1.5),
        "edu": ParagraphStyle("CVEdu", fontName=fm["regular"], fontSize=8, leading=11, textColor=TEXT_FULL, spaceAfter=1.5),
        "lang": ParagraphStyle("CVLang", fontName=fm["regular"], fontSize=8, leading=11, textColor=TEXT_FULL, spaceAfter=1.5),
    }


def _link(text, url):
    if url:
        return f'<a href="{url}" color="black">{text}</a>'
    return text


def _get_cv_data(version):
    from .models import CVProfile
    try:
        profile = CVProfile.objects.prefetch_related(
            'experiences', 'skills', 'projects', 'education', 'languages'
        ).get(version=version)
    except CVProfile.DoesNotExist:
        return None

    def _project_dict(proj):
        return {
            "name": proj.name,
            "type": proj.project_type,
            "date": proj.date,
            "desc": proj.description,
            "bullets": proj.get_bullets_list(),
            "link": proj.link_text,
            "url": proj.link_url,
        }

    all_projects = list(profile.projects.all())
    commercial = [_project_dict(p) for p in all_projects if p.section == "commercial"]
    personal = [_project_dict(p) for p in all_projects if p.section != "commercial"]

    return {
        "name": profile.name,
        "title": profile.title,
        "bio": profile.bio,
        "contact": {
            "email": profile.email,
            "location": profile.location,
            "linkedin": profile.linkedin,
            "linkedin_url": profile.linkedin_url,
            "website": profile.website,
        },
        "experience": [
            {
                "role": exp.role,
                "company": exp.company,
                "company_url": exp.company_url,
                "date": exp.date,
                "location": exp.location,
                "bullets": exp.get_bullets_list(),
            }
            for exp in profile.experiences.all()
        ],
        "skills": [
            (skill.category, skill.items)
            for skill in profile.skills.all()
        ],
        "projects_commercial": commercial,
        "projects_personal": personal,
        "projects_personal_label": "WORKS" if version == "composer" else "PERSONAL PROJECTS",
        "education": [
            (edu.title, edu.date, edu.url)
            for edu in profile.education.all()
        ],
        "languages": [
            (lang.language, lang.level)
            for lang in profile.languages.all()
        ],
        "updated_at": profile.updated_at,
    }


def _build_pdf(data):
    buffer = io.BytesIO()
    S = _get_styles()

    MARGIN_LR = 16 * mm
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=MARGIN_LR,
        rightMargin=MARGIN_LR,
        topMargin=14 * mm,
        bottomMargin=10 * mm,
    )

    elements = []

    elements.append(Paragraph(data["name"], S["name"]))
    elements.append(Spacer(1, 1))
    elements.append(Paragraph(data["title"], S["title"]))
    elements.append(Spacer(1, 2))

    contact_parts = []
    if data["contact"]["email"]:
        contact_parts.append(data["contact"]["email"])
    if data["contact"]["location"]:
        contact_parts.append(data["contact"]["location"])
    if data["contact"]["linkedin"]:
        contact_parts.append(_link(data["contact"]["linkedin"], data["contact"]["linkedin_url"]))
    if data["contact"]["website"]:
        contact_parts.append(_link(data["contact"]["website"], data["contact"]["website"]))
    elements.append(Paragraph("  |  ".join(contact_parts), S["contact"]))

    if data.get("bio"):
        elements.append(Spacer(1, 6))
        elements.append(Paragraph(data["bio"], S["bio"]))
    else:
        elements.append(Spacer(1, 8))

    if data["experience"]:
        elements.append(Paragraph("PROFESSIONAL EXPERIENCE", S["section"]))
        elements.append(Spacer(1, 4))
        for i, exp in enumerate(data["experience"]):
            elements.append(Paragraph(f'{exp["role"]}  \u2014  {exp["date"]}', S["role"]))
            company_text = _link(exp["company"], exp.get("company_url"))
            loc_part = f', {exp["location"]}' if exp.get("location") else ''
            elements.append(Paragraph(f'{company_text}{loc_part}', S["meta"]))
            elements.append(Spacer(1, 2))
            for bullet in exp["bullets"]:
                elements.append(Paragraph(f"- {bullet}", S["bullet"]))
            elements.append(Spacer(1, 8 if i < len(data["experience"]) - 1 else 5))

    if data["skills"]:
        elements.append(Spacer(1, 2))
        elements.append(Paragraph("SKILLS", S["section"]))
        elements.append(Spacer(1, 3))
        for cat, items in data["skills"]:
            elements.append(Paragraph(f"<b>{cat}:</b>  {items}", S["skill"]))

    def _render_project_group(header, projects):
        elements.append(Spacer(1, 6))
        elements.append(Paragraph(header, S["section"]))
        elements.append(Spacer(1, 3))
        for proj in projects:
            link_part = ""
            if proj.get("link") and proj.get("url"):
                link_part = f'  \u2014  {_link(proj["link"], proj["url"])}'
            name_str = f'{proj["name"]} ({proj["type"]})  \u2014  {proj["date"]}{link_part}'
            elements.append(Paragraph(name_str, S["project_name"]))
            elements.append(Paragraph(proj["desc"], S["project_desc"]))
            for bullet in proj.get("bullets") or []:
                elements.append(Paragraph(f"- {bullet}", S["bullet"]))
            elements.append(Spacer(1, 4))

    if data["projects_commercial"]:
        _render_project_group("COMMERCIAL PROJECTS", data["projects_commercial"])
    if data["projects_personal"]:
        _render_project_group(data["projects_personal_label"], data["projects_personal"])

    if data["education"]:
        elements.append(Spacer(1, 2))
        elements.append(Paragraph("EDUCATION &amp; CERTIFICATIONS", S["section"]))
        elements.append(Spacer(1, 3))
        for title, date, url in data["education"]:
            title_text = _link(title, url)
            elements.append(Paragraph(f"{title_text}  \u2014  {date}", S["edu"]))

    if data["languages"]:
        elements.append(Spacer(1, 6))
        elements.append(Paragraph("LANGUAGES", S["section"]))
        elements.append(Spacer(1, 3))
        for lang, level in data["languages"]:
            elements.append(Paragraph(f"<b>{lang}:</b>  {level}", S["lang"]))

    doc.build(elements)
    return buffer.getvalue()


def _build_docx(data):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn

    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.0)

    font_name = "Courier New"
    secondary = RGBColor(0x33, 0x33, 0x33)
    black_c = RGBColor(0x00, 0x00, 0x00)

    def add_run(para, text, size=8.5, bold=False, color=black_c):
        run = para.add_run(text)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:cs'), font_name)
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = color
        return run

    def add_hyperlink(para, text, url, size=8.5, bold=False, color=black_c):
        part = para.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink = para._element.makeelement(qn('w:hyperlink'), {qn('r:id'): r_id})
        new_run = para._element.makeelement(qn('w:r'), {})
        rPr = para._element.makeelement(qn('w:rPr'), {})
        rFonts = para._element.makeelement(qn('w:rFonts'), {qn('w:ascii'): font_name, qn('w:hAnsi'): font_name, qn('w:cs'): font_name})
        sz = para._element.makeelement(qn('w:sz'), {qn('w:val'): str(int(size * 2))})
        szCs = para._element.makeelement(qn('w:szCs'), {qn('w:val'): str(int(size * 2))})
        color_hex = str(color) if color else '000000'
        c = para._element.makeelement(qn('w:color'), {qn('w:val'): color_hex})
        rPr.append(rFonts)
        rPr.append(sz)
        rPr.append(szCs)
        rPr.append(c)
        if bold:
            b = para._element.makeelement(qn('w:b'), {})
            rPr.append(b)
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        para._element.append(hyperlink)

    p = doc.add_paragraph()
    p.space_after = Pt(0)
    p.space_before = Pt(0)
    add_run(p, data["name"], size=17, bold=True)

    p = doc.add_paragraph()
    p.space_after = Pt(1)
    p.space_before = Pt(0)
    add_run(p, data["title"], size=9, color=secondary)

    p = doc.add_paragraph()
    p.space_after = Pt(4)
    p.space_before = Pt(1)
    contact_parts = []
    if data["contact"]["email"]:
        contact_parts.append(data["contact"]["email"])
    if data["contact"]["location"]:
        contact_parts.append(data["contact"]["location"])
    add_run(p, "  |  ".join(contact_parts) + "  |  ", size=8.5, color=secondary)
    if data["contact"]["linkedin"] and data["contact"]["linkedin_url"]:
        add_hyperlink(p, data["contact"]["linkedin"], data["contact"]["linkedin_url"], size=8.5, color=secondary)

    if data.get("bio"):
        p = doc.add_paragraph()
        p.space_after = Pt(4)
        p.space_before = Pt(2)
        add_run(p, data["bio"], size=8.5, color=secondary)

    if data["experience"]:
        p = doc.add_paragraph()
        p.space_after = Pt(2)
        p.space_before = Pt(2)
        add_run(p, "PROFESSIONAL EXPERIENCE", size=8.5, bold=True, color=secondary)

        for exp in data["experience"]:
            p = doc.add_paragraph()
            p.space_after = Pt(0)
            p.space_before = Pt(2)
            add_run(p, f'{exp["role"]}  \u2014  {exp["date"]}', size=9.5, bold=True)

            p = doc.add_paragraph()
            p.space_after = Pt(1)
            p.space_before = Pt(0)
            if exp.get("company_url"):
                add_hyperlink(p, exp["company"], exp["company_url"], size=8.5, bold=True)
            else:
                add_run(p, exp["company"], size=8.5, bold=True)
            if exp.get("location"):
                add_run(p, f', {exp["location"]}', size=8.5, color=secondary)

            for bullet in exp["bullets"]:
                p = doc.add_paragraph()
                p.space_after = Pt(0)
                p.space_before = Pt(0)
                add_run(p, f"- {bullet}", size=8.5, color=secondary)

    if data["skills"]:
        p = doc.add_paragraph()
        p.space_after = Pt(2)
        p.space_before = Pt(3)
        add_run(p, "SKILLS", size=8.5, bold=True, color=secondary)

        for cat, items in data["skills"]:
            p = doc.add_paragraph()
            p.space_after = Pt(0)
            p.space_before = Pt(0)
            add_run(p, f"{cat}: ", size=8.5, bold=True, color=black_c)
            add_run(p, items, size=8.5, color=secondary)

    def _render_project_group(header, projects):
        p = doc.add_paragraph()
        p.space_after = Pt(2)
        p.space_before = Pt(3)
        add_run(p, header, size=8.5, bold=True, color=secondary)

        for proj in projects:
            p = doc.add_paragraph()
            p.space_after = Pt(0)
            p.space_before = Pt(1)
            add_run(p, f'{proj["name"]} ({proj["type"]})  \u2014  {proj["date"]}', size=8.5, bold=True)

            p = doc.add_paragraph()
            p.space_after = Pt(0)
            p.space_before = Pt(0)
            add_run(p, proj["desc"], size=8.5, color=secondary)
            if proj.get("link") and proj.get("url"):
                add_run(p, f'  \u2014  ', size=8.5, color=secondary)
                add_hyperlink(p, proj["link"], proj["url"], size=8.5, color=secondary)

            for bullet in proj.get("bullets") or []:
                p = doc.add_paragraph()
                p.space_after = Pt(0)
                p.space_before = Pt(0)
                add_run(p, f"- {bullet}", size=8.5, color=secondary)

    if data["projects_commercial"]:
        _render_project_group("COMMERCIAL PROJECTS", data["projects_commercial"])
    if data["projects_personal"]:
        _render_project_group(data["projects_personal_label"], data["projects_personal"])

    if data["education"]:
        p = doc.add_paragraph()
        p.space_after = Pt(2)
        p.space_before = Pt(3)
        add_run(p, "EDUCATION & CERTIFICATIONS", size=8.5, bold=True, color=secondary)

        for title, date, url in data["education"]:
            p = doc.add_paragraph()
            p.space_after = Pt(0)
            p.space_before = Pt(0)
            if url:
                add_hyperlink(p, title, url, size=8.5, bold=True)
            else:
                add_run(p, title, size=8.5, bold=True)
            add_run(p, f'  \u2014  {date}', size=8.5, color=secondary)

    if data["languages"]:
        p = doc.add_paragraph()
        p.space_after = Pt(2)
        p.space_before = Pt(3)
        add_run(p, "LANGUAGES", size=8.5, bold=True, color=secondary)

        for lang, level in data["languages"]:
            p = doc.add_paragraph()
            p.space_after = Pt(0)
            p.space_before = Pt(0)
            add_run(p, f"{lang}: ", size=8.5, bold=True)
            add_run(p, level, size=8.5, color=secondary)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _build_txt(data):
    lines = []
    lines.append(data["name"])
    lines.append(data["title"])
    lines.append("")

    contact_parts = []
    if data["contact"]["email"]:
        contact_parts.append(data["contact"]["email"])
    if data["contact"]["location"]:
        contact_parts.append(data["contact"]["location"])
    if data["contact"]["linkedin_url"]:
        contact_parts.append(data["contact"]["linkedin_url"])
    if data["contact"]["website"]:
        contact_parts.append(data["contact"]["website"])
    lines.append("  |  ".join(contact_parts))

    if data.get("bio"):
        lines.append("")
        lines.append(data["bio"])

    lines.append("")
    lines.append("")

    if data["experience"]:
        lines.append("PROFESSIONAL EXPERIENCE")
        lines.append("")
        for exp in data["experience"]:
            lines.append(f'{exp["role"]}  \u2014  {exp["date"]}')
            company_line = exp["company"]
            if exp.get("company_url"):
                company_line += f'  ({exp["company_url"]})'
            if exp.get("location"):
                company_line += f'  \u2014  {exp["location"]}'
            lines.append(company_line)
            for bullet in exp["bullets"]:
                lines.append(f"- {bullet}")
            lines.append("")

    if data["skills"]:
        lines.append("")
        lines.append("SKILLS")
        lines.append("")
        for cat, items in data["skills"]:
            lines.append(f"{cat}: {items}")

    def _render_project_group(header, projects):
        lines.append("")
        lines.append("")
        lines.append(header)
        lines.append("")
        for proj in projects:
            lines.append(f'{proj["name"]} \u00b7 {proj["type"]}  \u2014  {proj["date"]}')
            lines.append(proj["desc"])
            for bullet in proj.get("bullets") or []:
                lines.append(f"- {bullet}")
            if proj.get("url"):
                lines.append(proj["url"])
            lines.append("")

    if data["projects_commercial"]:
        _render_project_group("COMMERCIAL PROJECTS", data["projects_commercial"])
    if data["projects_personal"]:
        _render_project_group(data["projects_personal_label"], data["projects_personal"])

    if data["education"]:
        lines.append("")
        lines.append("EDUCATION & CERTIFICATIONS")
        lines.append("")
        for title, date, url in data["education"]:
            line = f"{title}  \u2014  {date}"
            if url:
                line += f"  ({url})"
            lines.append(line)

    if data["languages"]:
        lines.append("")
        lines.append("")
        lines.append("LANGUAGES")
        lines.append("")
        for lang, level in data["languages"]:
            lines.append(f"{lang}: {level}")

    lines.append("")
    return "\n".join(lines)


def _cache_path(version, fmt):
    return CACHE_DIR / f"cv_{version}.{fmt}"


def invalidate_cv_cache(version=None):
    versions = [version] if version else ['engineer', 'composer']
    for v in versions:
        for fmt in ('pdf', 'docx', 'txt'):
            path = _cache_path(v, fmt)
            if path.exists():
                path.unlink()


def get_cv_pdf(version='engineer'):
    path = _cache_path(version, 'pdf')
    if path.exists():
        return path.read_bytes()

    data = _get_cv_data(version)
    if not data:
        return None

    pdf_bytes = _build_pdf(data)
    path.write_bytes(pdf_bytes)
    return pdf_bytes


def get_cv_docx(version='engineer'):
    path = _cache_path(version, 'docx')
    if path.exists():
        return path.read_bytes()

    data = _get_cv_data(version)
    if not data:
        return None

    docx_bytes = _build_docx(data)
    path.write_bytes(docx_bytes)
    return docx_bytes


def get_cv_txt(version='engineer'):
    path = _cache_path(version, 'txt')
    if path.exists():
        return path.read_text(encoding='utf-8')

    data = _get_cv_data(version)
    if not data:
        return None

    txt_content = _build_txt(data)
    path.write_text(txt_content, encoding='utf-8')
    return txt_content
